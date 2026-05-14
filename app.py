import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
import io
import json

from database import (
    init_db,
    get_next_number,
    save_bill,
    get_all_bills,
    search_bill
)

# -----------------------------------
# INIT DATABASE
# -----------------------------------

init_db()

# -----------------------------------
# PAGE CONFIG
# -----------------------------------

st.set_page_config(
    page_title="Focus Marketing Billing Software",
    layout="wide"
)

st.title("Focus Marketing Solutions")

# -----------------------------------
# SESSION STATE
# -----------------------------------

if "invoice_items" not in st.session_state:
    st.session_state.invoice_items = []

# -----------------------------------
# SIDEBAR
# -----------------------------------

menu = st.sidebar.selectbox(
    "Menu",
    [
        "Create Document",
        "Search Bill",
        "View All Bills"
    ]
)

# -----------------------------------
# CREATE DOCUMENT
# -----------------------------------

if menu == "Create Document":

    document_type = st.selectbox(
        "Select Document Type",
        [
            "Invoice",
            "Salary Voucher"
        ]
    )

    document_number = get_next_number(
        document_type
    )

    # -----------------------------------
    # DATE
    # -----------------------------------

    invoice_date = st.date_input(
        "Select Bill Date"
    )

    # -----------------------------------
    # PAYMENT MODE
    # -----------------------------------

    payment_mode = st.selectbox(
        "Payment Mode",
        [
            "Cash",
            "UPI",
            "Bank Transfer",
            "Card",
            "Online"
        ]
    )

    transaction_id = ""

    if payment_mode != "Cash":

        transaction_id = st.text_input(
            "Transaction ID"
        )

    # ===================================
    # INVOICE
    # ===================================

    if document_type == "Invoice":

        st.header("Invoice Details")

        client_name = st.text_input(
            "Client Name"
        )

        company_name = st.text_input(
            "Company Name"
        )

        client_address = st.text_area(
            "Client Address"
        )

        client_gst = st.text_input(
            "Client GST Number"
        )

        gst_percentage = st.number_input(
            "GST %",
            min_value=0.0,
            value=18.0
        )

        remarks = st.text_area(
            "Remarks"
        )

        st.subheader("Services")

        with st.form("service_form"):

            col1, col2, col3 = st.columns(3)

            with col1:
                description = st.text_input(
                    "Service Description"
                )

            with col2:
                quantity = st.number_input(
                    "Quantity",
                    min_value=1,
                    value=1
                )

            with col3:
                rate = st.number_input(
                    "Rate",
                    min_value=0.0,
                    value=0.0
                )

            add_service = st.form_submit_button(
                "Add Service"
            )

            if add_service:

                amount = quantity * rate

                st.session_state.invoice_items.append({
                    "Description": description,
                    "Quantity": quantity,
                    "Rate": rate,
                    "Amount": amount
                })

        if len(st.session_state.invoice_items) > 0:

            df = pd.DataFrame(
                st.session_state.invoice_items
            )

            st.dataframe(df)

            # REMOVE OPTION

            remove_options = [
                f"{idx + 1}. {item['Description']}"
                for idx, item in enumerate(
                    st.session_state.invoice_items
                )
            ]

            selected_item = st.selectbox(
                "Remove Service",
                remove_options
            )

            if st.button("Remove Selected Service"):

                selected_index = remove_options.index(
                    selected_item
                )

                st.session_state.invoice_items.pop(
                    selected_index
                )

                st.rerun()

            subtotal = df["Amount"].sum()

            gst_amount = subtotal * (
                gst_percentage / 100
            )

            total_amount = subtotal + gst_amount

            st.metric(
                "Final Amount",
                f"₹ {total_amount:,.2f}"
            )

            # GENERATE DOC

            if st.button("Generate Invoice DOCX"):

                doc = Document()

                # HEADER TABLE

                table = doc.add_table(
                    rows=1,
                    cols=2
                )

                left = table.rows[0].cells[0]
                right = table.rows[0].cells[1]

                left.text = (
                    "FOCUS MARKETING SOLUTIONS"
                )

                right.text = (
                    f"Date: {invoice_date}"
                )

                doc.add_heading(
                    "INVOICE",
                    level=1
                )

                doc.add_paragraph(
                    f"Invoice Number: {document_number}"
                )

                doc.add_paragraph(
                    "GSTIN: __________________"
                )

                doc.add_heading(
                    "Client Details",
                    level=2
                )

                doc.add_paragraph(
                    f"Client Name: {client_name}"
                )

                doc.add_paragraph(
                    f"Company Name: {company_name}"
                )

                doc.add_paragraph(
                    f"Address: {client_address}"
                )

                doc.add_paragraph(
                    f"Client GST: {client_gst}"
                )

                # SERVICES TABLE

                service_table = doc.add_table(
                    rows=1,
                    cols=5
                )

                service_table.style = "Table Grid"

                hdr = service_table.rows[0].cells

                hdr[0].text = "Sl No"
                hdr[1].text = "Description"
                hdr[2].text = "Qty"
                hdr[3].text = "Rate"
                hdr[4].text = "Amount"

                for idx, item in enumerate(
                    st.session_state.invoice_items,
                    start=1
                ):

                    row = service_table.add_row().cells

                    row[0].text = str(idx)
                    row[1].text = item["Description"]
                    row[2].text = str(item["Quantity"])
                    row[3].text = str(item["Rate"])
                    row[4].text = str(item["Amount"])

                doc.add_paragraph(
                    f"Subtotal: ₹ {subtotal:,.2f}"
                )

                doc.add_paragraph(
                    f"GST ({gst_percentage}%): ₹ {gst_amount:,.2f}"
                )

                doc.add_paragraph(
                    f"Total Amount: ₹ {total_amount:,.2f}"
                )

                doc.add_paragraph(
                    f"Payment Mode: {payment_mode}"
                )

                if transaction_id != "":
                    doc.add_paragraph(
                        f"Transaction ID: {transaction_id}"
                    )

                doc.add_paragraph(
                    f"Remarks: {remarks}"
                )

                doc.add_paragraph(
                    "\nAuthorized Signature"
                )

                # SAVE TO DB

                save_bill(
                    document_number,
                    document_type,
                    client_name,
                    str(invoice_date),
                    payment_mode,
                    transaction_id,
                    total_amount,
                    json.dumps(
                        st.session_state.invoice_items
                    )
                )

                # DOWNLOAD

                buffer = io.BytesIO()

                doc.save(buffer)

                buffer.seek(0)

                st.download_button(
                    "Download Invoice",
                    data=buffer,
                    file_name=f"{document_number}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # ===================================
    # SALARY VOUCHER
    # ===================================

    else:

        st.header("Salary Voucher")

        employee_name = st.text_input(
            "Employee Name"
        )

        aadhaar_number = st.text_input(
            "Aadhaar Number"
        )

        pan_number = st.text_input(
            "PAN Number"
        )

        salary_amount = st.number_input(
            "Salary Amount",
            min_value=0.0
        )

        remarks = st.text_area(
            "Remarks"
        )

        if st.button(
            "Generate Salary Voucher"
        ):

            doc = Document()

            table = doc.add_table(
                rows=1,
                cols=2
            )

            left = table.rows[0].cells[0]
            right = table.rows[0].cells[1]

            left.text = (
                "FOCUS MARKETING SOLUTIONS"
            )

            right.text = (
                f"Date: {invoice_date}"
            )

            doc.add_heading(
                "SALARY VOUCHER",
                level=1
            )

            doc.add_paragraph(
                f"Voucher Number: {document_number}"
            )

            doc.add_paragraph(
                f"Employee Name: {employee_name}"
            )

            doc.add_paragraph(
                f"Aadhaar Number: {aadhaar_number}"
            )

            doc.add_paragraph(
                f"PAN Number: {pan_number}"
            )

            doc.add_paragraph(
                f"Salary Amount: ₹ {salary_amount:,.2f}"
            )

            doc.add_paragraph(
                f"Payment Mode: {payment_mode}"
            )

            if transaction_id != "":
                doc.add_paragraph(
                    f"Transaction ID: {transaction_id}"
                )

            doc.add_paragraph(
                f"Remarks: {remarks}"
            )

            doc.add_paragraph(
                "\nAuthorized Signature"
            )

            save_bill(
                document_number,
                document_type,
                employee_name,
                str(invoice_date),
                payment_mode,
                transaction_id,
                salary_amount,
                "{}"
            )

            buffer = io.BytesIO()

            doc.save(buffer)

            buffer.seek(0)

            st.download_button(
                "Download Salary Voucher",
                data=buffer,
                file_name=f"{document_number}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# -----------------------------------
# SEARCH BILL
# -----------------------------------

elif menu == "Search Bill":

    st.header("Search Bill")

    search_number = st.text_input(
        "Enter Invoice/Voucher Number"
    )

    if st.button("Search"):

        result = search_bill(search_number)

        if result:

            st.success("Bill Found")

            st.write(
                f"Document Number: {result[1]}"
            )

            st.write(
                f"Document Type: {result[2]}"
            )

            st.write(
                f"Customer Name: {result[3]}"
            )

            st.write(
                f"Date: {result[4]}"
            )

            st.write(
                f"Payment Mode: {result[5]}"
            )

            st.write(
                f"Amount: ₹ {result[7]}"
            )

        else:

            st.error("Bill Not Found")

# -----------------------------------
# VIEW ALL BILLS
# -----------------------------------

elif menu == "View All Bills":

    st.header("All Bills")

    rows = get_all_bills()

    if rows:

        df = pd.DataFrame(
            rows,
            columns=[
                "Document Number",
                "Document Type",
                "Customer Name",
                "Date",
                "Payment Mode",
                "Amount"
            ]
        )

        st.dataframe(
            df,
            use_container_width=True
        )

    else:

        st.info("No bills available")
